# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

doc = Document()

# ============ 全局样式设置 ============
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5

# 设置页边距
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

def add_heading_custom(text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        if level == 1:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0, 0, 0)
        elif level == 2:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 0, 0)
        elif level == 3:
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0, 0, 0)
    return heading

def add_para(text, bold=False, align=None, indent=True, font_size=12, font_name='宋体'):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.5
    return p

def add_bullet(text, level=0, font_size=12):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Pt(24 + level * 24)
    for run in p.runs:
        run.font.size = Pt(font_size)
    return p

def add_code_block(code_text):
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.0
    return p

def set_cell_shading(cell, color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.name = '黑体'
                run.font.size = Pt(11)
                run.bold = True
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        set_cell_shading(cell, "4472C4")
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)

    # Rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = '宋体'
                    run.font.size = Pt(10)
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            if r_idx % 2 == 0:
                set_cell_shading(cell, "D6E4F0")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table

# ============ 封面 ============
for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('毕业设计说明书')
run.font.name = '黑体'
run.font.size = Pt(36)
run.bold = True
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('校园社区综合服务与报修工单平台\n——EduFix')
run.font.name = '黑体'
run.font.size = Pt(22)
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

for _ in range(4):
    doc.add_paragraph()

info_items = [
    ('学    院：', '信息与计算机学院'),
    ('专    业：', '计算机科学与技术'),
    ('学生姓名：', '___________'),
    ('学    号：', '___________'),
    ('指导教师：', '___________'),
    ('完成日期：', datetime.datetime.now().strftime('%Y年%m月%d日')),
]

for label, value in info_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'{label}{value}')
    run.font.name = '宋体'
    run.font.size = Pt(16)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ============ 分页 - 摘要 ============
doc.add_page_break()

add_heading_custom('摘  要', level=1)

add_para('随着高校信息化建设的不断推进，校园后勤管理的数字化转型已成为提升服务质量和管理效率的关键。传统校园报修服务多依赖电话或线下登记，存在信息不透明、流程不可追溯、响应不及时等问题。为解决上述问题，本毕业设计设计并实现了一个名为"EduFix"的校园社区综合服务与报修工单平台。')

add_para('本系统采用前后端分离架构，后端基于 Spring Boot 2.7 框架搭建，使用 MyBatis-Plus 作为持久层框架，MySQL 作为关系型数据库，Redis 实现数据缓存和优先级队列；前端基于 Vue 3 框架和 Element Plus 组件库构建，使用 Pinia 进行状态管理，Vite 作为构建工具。系统实现了多角色权限管理（普通用户、维修员、管理员）、工单全生命周期管理（创建、派单、接单、处理、完成、评价）、公告发布与管理、服务评价体系等功能。')

add_para('本系统的特色在于充分利用 Redis 的数据结构特性，使用 Set 管理工单状态集合、ZSet 实现基于紧急程度的智能派单优先级队列、Hash 缓存高频访问数据，显著提升了系统性能与用户体验。系统工单状态流转清晰，权限控制完善，交互体验流畅，能够有效满足校园社区综合服务的实际需求。')

add_para('关键词：校园服务；工单管理；Spring Boot；Vue 3；Redis；前后端分离')

# ============ 分页 - Abstract ============
doc.add_page_break()

add_heading_custom('Abstract', level=1)

add_para('With the continuous advancement of information technology in universities, the digital transformation of campus logistics management has become a key factor in improving service quality and management efficiency. Traditional campus repair services mostly rely on phone calls or offline registration, which suffer from issues such as lack of transparency, untraceable processes, and slow response times. To address these problems, this graduation project designs and implements a campus comprehensive service and repair work order platform called "EduFix".')

add_para('The system adopts a front-end and back-end separation architecture. The back-end is built on the Spring Boot 2.7 framework, using MyBatis-Plus as the persistence layer framework, MySQL as the relational database, and Redis for data caching and priority queues. The front-end is built on the Vue 3 framework and Element Plus component library, using Pinia for state management and Vite as the build tool. The system implements multi-role permission management (regular users, maintenance staff, administrators), full lifecycle work order management (creation, assignment, acceptance, processing, completion, evaluation), announcement management, and service evaluation system.')

add_para('The distinctive feature of this system lies in its full utilization of Redis data structures: using Set to manage work order status collections, ZSet to implement an intelligent dispatch priority queue based on urgency levels, and Hash to cache frequently accessed data, significantly improving system performance and user experience. The system features clear work order status transitions, comprehensive permission control, and smooth interactive experience, effectively meeting the actual needs of campus community comprehensive services.')

add_para('Keywords: Campus Service; Work Order Management; Spring Boot; Vue 3; Redis; Front-end and Back-end Separation')

# ============ 目录页 ============
doc.add_page_break()

add_heading_custom('目  录', level=1)

toc_items = [
    ('摘要', 'I'),
    ('Abstract', 'II'),
    ('第一章 绪论', '1'),
    ('  1.1 项目背景', '1'),
    ('  1.2 项目目标与意义', '1'),
    ('  1.3 技术选型', '2'),
    ('第二章 系统需求分析', '3'),
    ('  2.1 功能需求', '3'),
    ('  2.2 非功能需求', '4'),
    ('  2.3 角色分析', '4'),
    ('第三章 系统设计', '5'),
    ('  3.1 系统架构设计', '5'),
    ('  3.2 数据库设计', '6'),
    ('  3.3 接口设计', '8'),
    ('  3.4 Redis数据结构设计', '9'),
    ('  3.5 工单状态机设计', '10'),
    ('第四章 系统实现', '11'),
    ('  4.1 后端实现', '11'),
    ('  4.2 前端实现', '13'),
    ('  4.3 核心功能实现', '15'),
    ('  4.4 Redis应用实现', '17'),
    ('第五章 系统测试', '18'),
    ('  5.1 测试环境', '18'),
    ('  5.2 功能测试', '18'),
    ('  5.3 接口测试', '19'),
    ('第六章 总结与展望', '20'),
    ('  6.1 项目总结', '20'),
    ('  6.2 技术亮点', '20'),
    ('  6.3 不足与展望', '21'),
    ('参考文献', '22'),
    ('致谢', '23'),
]

for item, page in toc_items:
    p = doc.add_paragraph()
    if item.startswith('  '):
        p.paragraph_format.left_indent = Cm(1.5)
        item = item.strip()
    run = p.add_run(item)
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    tab_run = p.add_run('\t' + page)
    tab_run.font.name = '宋体'
    tab_run.font.size = Pt(12)
    p.paragraph_format.line_spacing = 1.8

# ============ 第一章 绪论 ============
doc.add_page_break()
add_heading_custom('第一章 绪论', level=1)

add_heading_custom('1.1 项目背景', level=2)

add_para('随着高等教育事业的快速发展，高校校园规模不断扩大，师生人数持续增加，校园基础设施的维护和管理面临着越来越大的挑战。传统的校园报修模式通常采用电话报修、线下填写报修单等方式，这些方式存在以下突出问题：')

add_bullet('报修信息分散，难以集中管理和统计分析')
add_bullet('维修进度不透明，用户无法实时跟踪工单状态')
add_bullet('派单依赖人工经验，效率低下且缺乏公平性')
add_bullet('服务评价机制缺失，服务质量难以保证')
add_bullet('信息孤岛严重，各业务流程之间缺乏联动')

add_para('近年来，国家大力推进教育信息化建设，教育部明确提出要"推动信息技术与教育深度融合"。在此背景下，开发一套校园综合服务与报修工单管理平台，实现报修服务的数字化、智能化管理，具有重要的现实意义和应用价值。')

add_heading_custom('1.2 项目目标与意义', level=2)

add_para('本项目旨在设计并实现一个功能完善、体验良好的校园社区综合服务与报修工单平台（EduFix），具体目标包括：')

add_bullet('建立统一的工单管理平台，实现报修、建议、失物招领等综合服务的线上化')
add_bullet('实现多角色协同工作机制，支持普通用户、维修员、管理员三种角色的协同作业')
add_bullet('设计清晰的工单状态流转机制，实现工单全生命周期的信息化管理')
add_bullet('利用Redis实现智能派单和高效缓存，提升系统响应性能')
add_bullet('建立服务评价体系，形成服务质量反馈闭环')

add_para('本项目的意义在于：第一，通过信息化手段提升校园后勤服务效率和质量；第二，为高校信息化建设提供一种可借鉴的技术方案；第三，综合运用当前主流的后端开发技术（Spring Boot、Redis、MySQL）和前端开发技术（Vue 3、Element Plus），是一次完整的全栈开发实践。')

add_heading_custom('1.3 技术选型', level=2)

add_para('本系统在技术选型上坚持"成熟稳定、性能优异、生态丰富"的原则，具体技术栈如下：')

add_para('后端技术：')

add_table(
    ['技术/框架', '版本', '说明'],
    [
        ['Spring Boot', '2.7.14', 'Web框架与依赖注入'],
        ['MyBatis-Plus', '3.5.3.1', 'ORM持久层框架'],
        ['MySQL', '8.0.33', '关系型数据库'],
        ['Redis', '5.0+', '缓存与优先级队列'],
        ['JWT (jjwt)', '0.9.1', 'JSON Web Token认证'],
        ['Maven', '3.6+', '项目构建工具'],
        ['Lombok', '-', '减少样板代码'],
    ],
    col_widths=[4, 3, 8]
)

doc.add_paragraph()

add_para('前端技术：')

add_table(
    ['技术/框架', '版本', '说明'],
    [
        ['Vue 3', '3.5.34', '渐进式JavaScript框架'],
        ['Vite', '8.0.12', '构建工具与开发服务器'],
        ['Vue Router 4', '4.6.4', '前端路由管理'],
        ['Pinia', '3.0.4', '状态管理库'],
        ['Element Plus', '2.14.0', 'UI组件库'],
        ['Axios', '1.16.1', 'HTTP客户端'],
    ],
    col_widths=[4, 3, 8]
)

# ============ 第二章 系统需求分析 ============
doc.add_page_break()
add_heading_custom('第二章 系统需求分析', level=1)

add_heading_custom('2.1 功能需求', level=2)

add_para('本系统面向校园场景，提供综合服务与报修工单管理功能，主要包括以下功能模块：')

add_para('（1）用户认证模块', bold=True)
add_bullet('用户注册：新用户可以通过注册页面创建账号')
add_bullet('用户登录：支持用户名密码登录，返回JWT Token')
add_bullet('用户登出：将Token加入Redis黑名单，实现安全登出')

add_para('（2）工单管理模块', bold=True)
add_bullet('创建工单：支持报修（REPAIR）、建议（SUGGESTION）、失物招领（LOST_FOUND）三种类型')
add_bullet('工单列表：支持分页查询、按状态和类型筛选')
add_bullet('工单详情：展示工单完整信息、处理进度和操作历史')
add_bullet('工单取消：允许用户在特定状态下取消工单')
add_bullet('工单日志：记录工单的每一次状态变更')

add_para('（3）管理员功能模块', bold=True)
add_bullet('工单派单：管理员将待处理工单分配给维修员')
add_bullet('待处理工单查看：实时查看所有待处理的工单')
add_bullet('高优先级工单：基于Redis ZSet实现的优先级队列，按紧急程度排序')
add_bullet('员工管理：查看所有维修员列表及其信息')
add_bullet('公告管理：创建和发布系统公告')

add_para('（4）维修员功能模块', bold=True)
add_bullet('接单：维修员接受分配给自己的工单')
add_bullet('处理工单：开始处理并完成工单')
add_bullet('我的工单：查看分配给自己的工单列表')

add_para('（5）服务评价模块', bold=True)
add_bullet('创建评价：工单完成后，用户可以对服务进行评分（1-5星）、选择标签、填写评价内容')
add_bullet('查看评价：查看指定工单或维修员的评价信息')

add_para('（6）公告模块', bold=True)
add_bullet('公告列表：展示已发布的公告')
add_bullet('公告详情：查看公告具体内容')
add_bullet('公告管理：管理员创建、发布公告，支持Redis缓存')

add_heading_custom('2.2 非功能需求', level=2)

add_para('（1）性能需求：系统响应时间应在合理范围内，列表查询响应时间不超过2秒；利用Redis缓存减少数据库查询次数。')
add_para('（2）安全性需求：使用JWT进行身份认证，密码采用BCrypt加密存储，接口权限校验严格。')
add_para('（3）可用性需求：系统应具有良好的用户体验，界面简洁直观，操作反馈及时。')
add_para('（4）可扩展性需求：系统架构应支持功能模块的横向扩展，便于后续功能迭代。')
add_para('（5）兼容性需求：前端应支持主流浏览器（Chrome、Firefox、Edge等）。')

add_heading_custom('2.3 角色分析', level=2)

add_para('本系统设计了三种用户角色，各角色的职责和权限如下：')

add_table(
    ['角色', '职责', '核心权限', '典型用户'],
    [
        ['普通用户(USER)', '提交工单、查看进度、评价服务', '创建工单、查看自己的工单、评价已完成工单、浏览公告', '在校学生、教职工'],
        ['维修员(STAFF)', '接单、处理工单、完成工单', '查看分配给自己的工单、接单、完成工单', '后勤维修人员'],
        ['管理员(ADMIN)', '派单、管理公告、系统管理', '查看所有工单、派单、发布公告、数据统计', '后勤管理人员'],
    ],
    col_widths=[3, 4, 5, 3]
)

# ============ 第三章 系统设计 ============
doc.add_page_break()
add_heading_custom('第三章 系统设计', level=1)

add_heading_custom('3.1 系统架构设计', level=2)

add_para('本系统采用前后端分离的架构模式。前端和后端通过RESTful API进行通信，使用JSON作为数据交换格式。整体架构分为以下几层：')

add_para('（1）前端展示层：基于Vue 3框架，采用组件化开发方式，通过Vue Router实现路由管理，Pinia进行全局状态管理。前端使用Axios封装HTTP请求，统一处理Token认证和错误响应。')

add_para('（2）后端服务层：基于Spring Boot框架，采用分层架构设计，包括Controller层（接收请求）、Service层（业务逻辑）、Mapper层（数据访问）。使用JWT拦截器实现接口权限控制。')

add_para('（3）数据持久层：使用MySQL存储业务数据，MyBatis-Plus作为ORM框架简化数据库操作。使用Redis作为缓存层，提升热点数据访问性能，并利用其丰富的数据结构实现特定业务功能。')

add_para('系统架构示意图如下：')

add_code_block('''
┌─────────────────────────────────────────────────────────┐
│                    前端（Vue 3 + Element Plus）          │
│  ┌──────────┐  ┌──────────┐  ┌──────────┐  ┌────────┐ │
│  │ 认证页面  │  │ 工单页面  │  │ 管理页面  │  │公告页面│ │
│  └────┬─────┘  └────┬─────┘  └────┬─────┘  └────┬───┘ │
│       └──────────────┴──────────────┴──────────────┘    │
│                         Axios                           │
│                    (JWT Token 拦截器)                      │
└────────────────────────┬────────────────────────────────┘
                         │ RESTful API (JSON)
┌────────────────────────┴────────────────────────────────┐
│                    后端（Spring Boot）                    │
│  ┌────────────────────────────────────────────────────┐ │
│  │              Controller 层                          │ │
│  │  Auth │ Ticket │ Admin │ Staff │ Evaluation│Notice │ │
│  └──────────────────────┬─────────────────────────────┘ │
│  ┌──────────────────────┴─────────────────────────────┐ │
│  │              Service 层（业务逻辑）                  │ │
│  └──────────────────────┬─────────────────────────────┘ │
│  ┌──────────────────────┴─────────────────────────────┐ │
│  │              Mapper 层（MyBatis-Plus）              │ │
│  └────────────────────────────────────────────────────┘ │
└────────────────┬────────────────────────────────────────┘
                 │
    ┌────────────┴────────────┐
    │                         │
┌───┴────┐             ┌─────┴────┐
│  MySQL  │             │  Redis   │
│ 主数据库 │             │  缓存    │
└────────┘             └──────────┘
''')

add_heading_custom('3.2 数据库设计', level=2)

add_para('本系统使用MySQL作为主数据库，设计了7张数据表，涵盖了用户、工单、评价、公告等核心业务数据。')

add_para('3.2.1 数据库整体结构', bold=True)

add_table(
    ['表名', '说明', '核心字段'],
    [
        ['user', '用户表', 'id, username, password, role, real_name, identifier_no'],
        ['staff', '员工表', 'id, user_id, employee_no, department, specialty, rating'],
        ['ticket', '工单表', 'id, ticket_no, title, description, category, status, urgency, user_id, staff_id, images'],
        ['ticket_log', '工单日志表', 'id, ticket_id, action, from_status, to_status, content, operator'],
        ['evaluation', '评价表', 'id, ticket_id, user_id, staff_id, rating, tags, content'],
        ['notice', '公告表', 'id, title, content, type, priority, status, publish_time'],
        ['notification', '通知表', 'id, user_id, title, content, type, is_read'],
    ],
    col_widths=[3, 4, 8]
)

doc.add_paragraph()

add_para('3.2.2 核心表结构说明', bold=True)

add_para('user表：存储所有用户信息，包含用户名、密码（BCrypt加密）、角色（USER/STAFF/ADMIN）、真实姓名、学号/工号等字段。角色字段决定了用户在系统中的权限范围。', indent=True)

add_para('ticket表：工单核心表，包含以下关键字段：\n• category（工单分类）：REPAIR（报修）、SUGGESTION（建议）、LOST_FOUND（失物招领）\n• status（工单状态）：PENDING（待处理）、ASSIGNED（已派单）、IN_PROGRESS（处理中）、COMPLETED（已完成）、CANCELLED（已取消）\n• urgency（紧急程度）：LOW（低）、NORMAL（普通）、HIGH（高）、URGENT（紧急）\n• images：存储工单图片附件路径，JSON数组格式', indent=True)

add_para('evaluation表：评价表与工单一一对应（唯一约束），包含评分（1-5分）、评价标签、评价内容等字段。评分可汇总至staff表的rating字段。', indent=True)

add_heading_custom('3.3 接口设计', level=2)

add_para('系统后端对外提供RESTful API接口，遵循统一的响应格式：')

add_code_block('''{
  "code": 200,      // 状态码（200成功，401未授权，500服务端错误）
  "message": "操作成功",
  "data": {}        // 响应数据
}''')

add_para('主要API接口分类如下：')

add_para('（1）认证接口（公开）', bold=True)
add_table(
    ['端点', '方法', '说明'],
    [
        ['/api/auth/login', 'POST', '用户登录，返回JWT Token'],
        ['/api/auth/register', 'POST', '用户注册'],
        ['/api/auth/logout', 'POST', '用户登出（需认证）'],
    ],
    col_widths=[5, 2, 8]
)

doc.add_paragraph()

add_para('（2）工单接口（需认证）', bold=True)
add_table(
    ['端点', '方法', '说明'],
    [
        ['/api/ticket', 'POST', '创建工单'],
        ['/api/ticket/{id}', 'GET', '获取工单详情'],
        ['/api/ticket/list', 'GET', '获取工单列表（支持分页筛选）'],
        ['/api/ticket/{id}/logs', 'GET', '获取工单日志'],
        ['/api/ticket/{id}/cancel', 'PUT', '取消工单'],
    ],
    col_widths=[5, 2, 8]
)

doc.add_paragraph()

add_para('（3）管理员接口（需ADMIN角色）', bold=True)
add_table(
    ['端点', '方法', '说明'],
    [
        ['/api/admin/assign', 'POST', '派单（分配维修员）'],
        ['/api/admin/pending-tickets', 'GET', '获取待处理工单列表'],
        ['/api/admin/high-priority-tickets', 'GET', '获取高优先级工单'],
        ['/api/admin/staff/list', 'GET', '获取员工列表'],
    ],
    col_widths=[5, 2, 8]
)

doc.add_paragraph()

add_para('（4）维修员接口（需STAFF角色）', bold=True)
add_table(
    ['端点', '方法', '说明'],
    [
        ['/api/staff/{ticketId}/accept', 'POST', '接单'],
        ['/api/staff/{ticketId}/complete', 'POST', '完成工单'],
        ['/api/staff/my-tickets', 'GET', '获取我的工单列表'],
    ],
    col_widths=[5, 2, 8]
)

doc.add_paragraph()

add_para('（5）评价接口', bold=True)
add_table(
    ['端点', '方法', '说明'],
    [
        ['/api/evaluation', 'POST', '创建评价'],
        ['/api/evaluation/ticket/{ticketId}', 'GET', '获取工单评价'],
        ['/api/evaluation/staff/{staffId}', 'GET', '获取员工评价列表'],
    ],
    col_widths=[5, 2, 8]
)

doc.add_paragraph()

add_para('（6）公告接口', bold=True)
add_table(
    ['端点', '方法', '说明'],
    [
        ['/api/notice/published', 'GET', '获取已发布公告（公开）'],
        ['/api/notice/{id}', 'GET', '获取公告详情（公开）'],
        ['/api/notice', 'POST', '创建公告（管理员）'],
        ['/api/notice/{id}/publish', 'PUT', '发布公告（管理员）'],
    ],
    col_widths=[5, 2, 8]
)

add_heading_custom('3.4 Redis数据结构设计', level=2)

add_para('本系统充分运用Redis的多种数据结构，在缓存加速和业务功能两个方面发挥重要作用。具体设计如下：')

add_table(
    ['Redis数据结构', 'Key模式', '用途说明', '应用场景'],
    [
        ['String', 'token:blacklist:{token}', 'JWT Token黑名单', '用户登出后将Token加入黑名单，有效期为Token剩余有效期'],
        ['String', 'notice:published:list', '公告列表缓存', '缓存已发布公告列表，减少数据库查询'],
        ['Hash', 'notice:detail:{id}', '公告详情缓存', '缓存单个公告详情，支持快速读取'],
        ['Hash', 'evaluation:tags', '评价标签统计', '统计各评价标签的使用次数'],
        ['Set', 'ticket:status:{PENDING\|...\|CANCELLED}', '工单状态集合', '按状态维护工单ID集合，支持高效的状态筛选和统计'],
        ['ZSet', 'ticket:priority:queue', '优先级派单队列', '按紧急程度+提交时间计算分值，分数越高优先级越高'],
    ],
    col_widths=[2.5, 5, 3.5, 4.5]
)

doc.add_paragraph()

add_para('其中，优先级队列的分数计算公式为：priorityScore = urgencyWeight × 1000000000 + (Long.MAX_VALUE - timestamp)，其中urgencyWeight根据紧急程度（LOW=1, NORMAL=2, HIGH=3, URGENT=4）确定。这种设计确保紧急程度高的工单优先被处理，同时保证同等紧急程度下先提交的工单优先。')

add_heading_custom('3.5 工单状态机设计', level=2)

add_para('工单状态机是本系统的核心业务逻辑之一，它定义了工单从创建到完成的完整生命周期。状态机设计如下：')

add_code_block('''
                    ┌──────────┐
                    │ PENDING  │  (待处理)
                    └────┬─────┘
                         │ 管理员派单
                    ┌────▼─────┐
                    │ ASSIGNED │  (已派单)
                    └────┬─────┘
                         │ 维修员接单
                   ┌─────▼──────┐
                   │ IN_PROGRESS│  (处理中)
                   └─────┬──────┘
                         │ 维修员完成
                   ┌─────▼──────┐
                   │ COMPLETED  │  (已完成)
                   └────────────┘

  取消操作可在 PENDING / ASSIGNED / IN_PROGRESS 任一状态下执行：
        ┌──────────┐    ┌──────────┐    ┌──────────┐
        │ PENDING  │───→│ ASSIGNED │───→│IN_PROGRESS│
        └────┬─────┘    └────┬─────┘    └────┬──────┘
             │取消            │取消            │取消
             └───────┬───────┴───────┬───────┘
                     │ CANCELLED    │
                     └──────────────┘
''')

add_para('状态说明：')
add_bullet('PENDING（待处理）：工单刚创建，等待管理员派单')
add_bullet('ASSIGNED（已派单）：管理员已分配维修员，等待维修员接单')
add_bullet('IN_PROGRESS（处理中）：维修员已接单，正在处理')
add_bullet('COMPLETED（已完成）：维修员完成处理，工单结束')
add_bullet('CANCELLED（已取消）：用户在PENDING/ASSIGNED/IN_PROGRESS状态下取消工单')

# ============ 第四章 系统实现 ============
doc.add_page_break()
add_heading_custom('第四章 系统实现', level=1)

add_heading_custom('4.1 后端实现', level=2)

add_para('4.1.1 项目结构', bold=True)

add_para('后端项目基于Maven构建，采用标准的Spring Boot分层结构，主要包含以下包：')

add_code_block('''
edufix/src/main/java/org/example/edufix/
├── common/          # 通用类（统一响应Result）
├── config/          # 配置类（CORS、Jackson、Redis、Security）
├── controller/      # 控制器层（接收HTTP请求）
├── dto/             # 数据传输对象
├── entity/          # 实体类（数据库映射）
├── exception/       # 异常处理（全局异常处理）
├── handler/         # MyBatis-Plus自动填充处理器
├── interceptor/     # JWT认证拦截器
├── mapper/          # MyBatis-Plus Mapper接口
├── service/         # 业务逻辑层
├── util/            # 工具类（JWT、Redis）
└── EdufixApplication.java  # 启动类
''')

add_para('4.1.2 JWT认证实现', bold=True)

add_para('系统使用JWT（JSON Web Token）实现无状态认证。用户登录成功后，后端生成包含用户ID、用户名、角色等信息的Token返回给前端。前端在后续请求中通过Authorization头携带Token。', indent=True)

add_para('认证流程如下：')
add_bullet('用户发送登录请求，后端验证用户名密码')
add_bullet('验证通过后，使用JJWT库生成Token，过期时间设置为24小时')
add_bullet('前端将Token存储到localStorage，并在每次请求时通过Axios拦截器自动添加Authorization头')
add_bullet('后端JwtInterceptor拦截除登录、注册、公告查询外的所有请求，解析并验证Token')
add_bullet('Token验证失败返回401状态码，前端收到401时自动跳转到登录页面')

add_para('4.1.3 工单核心业务实现', bold=True)

add_para('工单业务是系统的核心，TicketService实现了完整的工单生命周期管理。关键实现要点如下：', indent=True)

add_para('创建工单时，系统生成唯一的工单编号（ticket_no），格式为"TK" + 时间戳 + 随机数。工单创建后自动进入PENDING状态，并记录创建日志。同时，工单ID被添加到Redis的Set集合（ticket:status:PENDING）和ZSet优先级队列（ticket:priority:queue）中。', indent=True)

add_para('派单操作由管理员执行，系统将工单状态从PENDING变更为ASSIGNED，记录分配的维修员ID和派单日志。同时更新Redis中的状态集合和优先级队列。', indent=True)

add_para('接单和完成操作由维修员执行，接单时状态从ASSIGNED变为IN_PROGRESS，完成时从IN_PROGRESS变为COMPLETED。每次状态变更都会记录详细的工单日志，并同步更新Redis数据结构。', indent=True)

add_para('4.1.4 评价服务实现', bold=True)

add_para('评价服务（EvaluationService）实现了工单完成后的服务评价功能。用户可以在工单完成后对维修员的服务进行评分（1-5分）、选择标签和填写评价内容。评价使用Redis Hash记录标签统计数据（evaluation:tags），同时更新员工的平均评分（staff.rating）。', indent=True)

add_heading_custom('4.2 前端实现', level=2)

add_para('4.2.1 项目结构', bold=True)

add_para('前端项目基于Vue 3 + Vite构建，采用Composition API（<script setup>）开发模式，项目结构如下：')

add_code_block('''
edufix-web/src/
├── api/            # API接口模块（封装Axios请求）
│   ├── auth.js     # 认证API
│   ├── ticket.js   # 工单API
│   ├── admin.js    # 管理员API
│   ├── staff.js    # 维修员API
│   ├── evaluation.js  # 评价API
│   └── notice.js   # 公告API
├── layout/         # 布局组件（MainLayout）
├── router/         # 路由配置（含路由守卫）
├── stores/         # Pinia状态管理
├── styles/         # 全局样式
├── utils/          # 工具函数（Axios封装）
├── views/          # 页面组件（13个页面）
│   ├── Login.vue
│   ├── Register.vue
│   ├── Dashboard.vue
│   ├── ticket/     # 工单相关页面
│   ├── admin/      # 管理员页面
│   ├── staff/      # 维修员页面
│   └── notice/     # 公告页面
├── App.vue
└── main.js
''')

add_para('4.2.2 路由与权限控制', bold=True)

add_para('系统使用Vue Router 4实现前端路由管理，并利用路由守卫（beforeEach）实现基于角色的权限控制。路由配置中为每个路由指定了所需的角色（meta.role），路由守卫根据用户角色进行访问控制。未登录用户访问受保护页面时自动跳转到登录页，角色不匹配时跳转到对应的默认页面。')

add_para('主要路由配置如下：')

add_table(
    ['路由路径', '页面组件', '访问权限', '说明'],
    [
        ['/login', 'Login.vue', '公开', '登录页面'],
        ['/register', 'Register.vue', '公开', '注册页面'],
        ['/dashboard', 'Dashboard.vue', 'USER', '用户首页'],
        ['/tickets', 'TicketList.vue', 'USER', '工单列表'],
        ['/tickets/create', 'CreateTicket.vue', 'USER', '创建工单'],
        ['/tickets/:id', 'TicketDetail.vue', 'USER/STAFF', '工单详情'],
        ['/admin/dashboard', 'AdminDashboard.vue', 'ADMIN', '管理后台'],
        ['/admin/assign', 'AssignTicket.vue', 'ADMIN', '工单派单'],
        ['/admin/notices', 'ManageNotices.vue', 'ADMIN', '公告管理'],
        ['/staff/dashboard', 'StaffDashboard.vue', 'STAFF', '维修员工作台'],
        ['/notices', 'NoticeList.vue', '所有', '公告列表'],
        ['/notices/:id', 'NoticeDetail.vue', '所有', '公告详情'],
    ],
    col_widths=[3, 3, 3, 3.5]
)

doc.add_paragraph()

add_para('4.2.3 状态管理与API封装', bold=True)

add_para('系统使用Pinia进行全局状态管理，主要管理用户认证状态。user store中包含登录、登出、Token管理、用户信息持久化等功能。用户信息存储到localStorage，刷新页面后自动恢复。', indent=True)

add_para('API调用通过Axios封装（utils/request.js），实现了请求拦截器和响应拦截器。请求拦截器自动从localStorage获取Token并添加到Authorization头；响应拦截器统一处理错误响应，401状态码自动触发登出跳转。', indent=True)

add_heading_custom('4.3 核心功能实现', level=2)

add_para('4.3.1 用户登录与注册', bold=True)

add_para('登录页面采用渐变背景设计，表单包含用户名和密码输入框，支持表单验证。登录成功后，后端返回Token和用户信息，前端将其存储到Pinia store和localStorage，然后根据角色跳转到对应的首页。', indent=True)

add_para('注册页面包含用户名、密码、确认密码、真实姓名、手机号、邮箱等字段，所有字段均经过前端验证（格式校验、密码强度校验、确认密码一致性校验等）。注册成功后自动跳转到登录页。', indent=True)

add_para('4.3.2 工单创建与浏览', bold=True)

add_para('创建工单页面支持三种工单类型（报修、建议、失物招领），根据类型动态显示不同的表单字段。报修工单需要填写位置信息和联系方式，失物招领工单需要填写物品描述和丢失/拾取地点。支持设置紧急程度（低/普通/高/紧急）。', indent=True)

add_para('工单列表页面采用分页加载模式，支持按状态和类型筛选，以卡片或表格形式展示工单信息。工单详情页面展示完整的工单信息、处理进度时间线和评价表单。进度时间线使用Element Plus的Timeline组件，直观展示工单状态变更历史。', indent=True)

add_para('4.3.3 管理员派单', bold=True)

add_para('管理员派单页面展示所有待处理的工单，管理员可以选择合适的维修员进行派单。系统同时提供高优先级工单列表，基于Redis ZSet实现，按紧急程度排序，帮助管理员优先处理紧急工单。', indent=True)

add_para('4.3.4 维修员工作台', bold=True)

add_para('维修员工作台展示分配给自己的工单，提供接单和完成操作。接单后工单状态变为处理中（IN_PROGRESS），完成后变为已完成（COMPLETED）。工作台还展示工单统计信息，如待处理工单数、处理中工单数、本月完成数等。', indent=True)

add_para('4.3.5 服务评价', bold=True)

add_para('用户可以在工单完成后对维修员的服务进行评价。评价采用五星评分系统，支持选择评价标签（如"快速响应"、"专业认真"、"服务态度好"等）和填写文字评价内容。评价完成后，系统更新维修员的综合评分。', indent=True)

add_heading_custom('4.4 Redis应用实现', level=2)

add_para('本系统是Redis在校园服务管理场景中的一次典型应用实践，充分发挥了Redis多种数据结构的特性。具体实现如下：')

add_para('（1）Token黑名单（String）：用户登出时，将JWT Token存入Redis并设置与Token一致的过期时间。拦截器在验证Token时，先检查该Token是否在黑名单中，有效防止Token在有效期内被重复使用。')

add_para('（2）公告缓存（String + Hash）：系统将已发布的公告列表缓存为JSON字符串（notice:published:list），将公告详情缓存为Hash结构（notice:detail:{id}）。管理员发布新公告时同步更新缓存，保证数据一致性。')

add_para('（3）工单状态集合（Set）：为每种工单状态维护一个Set集合（ticket:status:PENDING等），存储对应状态的工单ID。创建工单时向对应集合添加ID，状态变更时从原集合移除并加入新集合。这一设计支持O(1)时间复杂度的状态查询和统计。')

add_para('（4）优先级队列（ZSet）：工单创建时，按照"紧急程度×时间戳"的复合分值加入ZSet。分值计算公式为：score = urgencyWeight × 1000000000 + (Long.MAX_VALUE - timestamp)。管理员获取高优先级工单时，通过ZREVRANGE命令直接获取排序后的工单列表。')

# ============ 第五章 系统测试 ============
doc.add_page_break()
add_heading_custom('第五章 系统测试', level=1)

add_heading_custom('5.1 测试环境', level=2)

add_table(
    ['环境', '配置'],
    [
        ['操作系统', 'Windows 11'],
        ['后端运行环境', 'JDK 8, Maven 3.6+'],
        ['数据库', 'MySQL 8.0.33'],
        ['缓存', 'Redis 5.0+'],
        ['前端运行环境', 'Node.js 18+, Vite 8.0'],
        ['浏览器', 'Chrome 最新版, Firefox 最新版, Edge 最新版'],
        ['测试工具', 'Postman, curl, 浏览器开发者工具'],
    ],
    col_widths=[4, 11]
)

add_heading_custom('5.2 功能测试', level=2)

add_para('系统功能测试覆盖了所有核心功能模块，测试结果如下：')

add_table(
    ['测试模块', '测试用例', '预期结果', '测试结果'],
    [
        ['用户认证', '用户注册', '注册成功，跳转登录页', '✅ 通过'],
        ['用户认证', '用户登录', '登录成功，返回Token', '✅ 通过'],
        ['用户认证', '用户登出', '登出成功，Token失效', '✅ 通过'],
        ['工单管理', '创建报修工单', '工单创建成功，状态为PENDING', '✅ 通过'],
        ['工单管理', '创建建议工单', '工单创建成功', '✅ 通过'],
        ['工单管理', '创建失物招领工单', '工单创建成功', '✅ 通过'],
        ['工单管理', '查询工单列表', '返回分页工单数据', '✅ 通过'],
        ['工单管理', '查看工单详情', '返回完整的工单信息', '✅ 通过'],
        ['工单管理', '取消工单', '工单状态变更为CANCELLED', '✅ 通过'],
        ['管理员功能', '派单', '工单状态变更为ASSIGNED', '✅ 通过'],
        ['管理员功能', '查看待处理工单', '返回待处理工单列表', '✅ 通过'],
        ['管理员功能', '查看高优先级工单', '按优先级排序返回', '✅ 通过'],
        ['管理员功能', '创建公告', '公告创建成功', '✅ 通过'],
        ['管理员功能', '发布公告', '公告状态变更为已发布', '✅ 通过'],
        ['维修员功能', '接单', '工单状态变更为IN_PROGRESS', '✅ 通过'],
        ['维修员功能', '完成工单', '工单状态变更为COMPLETED', '✅ 通过'],
        ['维修员功能', '查看我的工单', '只返回分配给当前维修员的工单', '✅ 通过'],
        ['服务评价', '创建评价', '评价创建成功', '✅ 通过'],
        ['服务评价', '查看工单评价', '返回评价内容', '✅ 通过'],
        ['服务评价', '查看员工评价列表', '返回员工的所有评价', '✅ 通过'],
    ],
    col_widths=[2.5, 3.5, 4.5, 1.5]
)

add_heading_custom('5.3 接口测试', level=2)

add_para('使用Postman和curl对所有API接口进行了测试，测试覆盖认证接口、工单接口、管理员接口、维修员接口、评价接口和公告接口。所有接口均能正确响应，返回统一格式的JSON数据。', indent=True)

add_para('测试验证了以下几点：\n• 公开接口可以直接访问，无需Token认证\n• 需要认证的接口在缺少Token时返回401状态码\n• 管理员接口在非管理员角色访问时返回权限错误\n• 参数校验正常工作，非法参数返回相应的错误提示\n• 工单状态流转严格按照状态机设计执行\n• Redis缓存数据与数据库数据保持一致', indent=True)

# ============ 第六章 总结与展望 ============
doc.add_page_break()
add_heading_custom('第六章 总结与展望', level=1)

add_heading_custom('6.1 项目总结', level=2)

add_para('本毕业设计项目"EduFix校园社区综合服务与报修工单平台"经过需求分析、系统设计、编码实现和测试验证，已顺利完成所有既定功能模块的开发。系统实现了以下核心功能：')

add_bullet('多角色权限管理系统：支持普通用户、维修员、管理员三种角色的协同工作')
add_bullet('工单全生命周期管理：覆盖创建、派单、接单、处理、完成、评价的完整流程')
add_bullet('智能派单机制：基于Redis ZSet实现按紧急程度排序的优先级派单队列')
add_bullet('Redis多场景应用：使用String、Hash、Set、ZSet四种数据结构，分别实现缓存加速和业务功能')
add_bullet('服务评价体系：支持五星评分、标签选择和文字评价')
add_bullet('公告发布系统：支持公告的创建、发布和缓存加速')
add_bullet('前后端分离架构：前端Vue 3 + Element Plus，后端Spring Boot + MyBatis-Plus')

add_para('在开发过程中，严格遵循软件工程规范，进行了系统化的需求分析、架构设计、编码实现和测试验证。项目代码结构清晰、注释完整、命名规范，具备良好的可读性和可维护性。')

add_heading_custom('6.2 技术亮点', level=2)

add_para('（1）Redis的深度应用：系统不仅将Redis用于常规的数据缓存，还创造性地利用Set集合管理工单状态、ZSet有序集合实现优先级派单队列，充分发挥了Redis多种数据结构的优势。')

add_para('（2）完善的工单状态机：设计的五状态工单状态机（PENDING→ASSIGNED→IN_PROGRESS→COMPLETED/CANCELLED）逻辑清晰、流转明确，每次状态变更均有详细日志记录，确保工单全流程可追溯。')

add_para('（3）前后端分离架构：采用当前业界主流的前后端分离架构，前端使用Vue 3 Composition API，后端使用Spring Boot分层架构，前后端通过RESTful API通信，职责清晰、易于扩展。')

add_para('（4）多角色权限控制：前端通过路由守卫实现页面级权限控制，后端通过JWT拦截器和角色校验实现接口级权限控制，形成双重安全保障。')

add_para('（5）良好的用户体验：前端使用Element Plus组件库构建现代化UI界面，页面带有过渡动画、加载状态和错误提示，交互反馈及时、用户体验流畅。')

add_heading_custom('6.3 不足与展望', level=2)

add_para('尽管本系统已实现了核心功能，但仍存在一些不足和可改进的方面：')

add_para('（1）实时通知：当前系统缺乏实时通知机制，用户无法即时获知工单状态变更。未来可以引入WebSocket技术，实现工单状态变化的实时推送。')

add_para('（2）数据可视化：管理后台的统计数据较为基础，可以引入ECharts等数据可视化库，实现工单量趋势图、维修员绩效对比图、满意度分析图等可视化报表。')

add_para('（3）文件上传：当前系统的图片上传功能较为基础，可以进一步支持多图片上传、文件类型限制、图片压缩等功能。')

add_para('（4）消息中心：可以扩展通知模块，实现站内信系统，支持管理员向特定用户或用户组发送通知。')

add_para('（5）移动端适配：当前系统虽然具有一定响应式能力，但尚未针对移动端进行专门优化。未来可以开发微信小程序或移动端适配版本。')

add_para('（6）性能优化：随着数据量的增长，可以引入更完善的缓存策略和数据库索引优化，进一步提升系统性能。')

# ============ 参考文献 ============
doc.add_page_break()
add_heading_custom('参考文献', level=1)

refs = [
    '[1] 陈刚. Spring Boot从入门到实战[M]. 北京: 清华大学出版社, 2020.',
    '[2] 黄勇. Vue.js 3.0从入门到精通[M]. 北京: 机械工业出版社, 2021.',
    '[3] 李刚. 疯狂Java讲义[M]. 北京: 电子工业出版社, 2018.',
    '[4] Josiah L. Carlson. Redis实战[M]. 北京: 人民邮电出版社, 2015.',
    '[5] 王振宇. MyBatis-Plus实战指南[M]. 北京: 机械工业出版社, 2022.',
    '[6] 李云. 前后端分离架构设计实践[J]. 计算机应用与软件, 2021, 38(5): 15-20.',
    '[7] Johnson R. Professional Java Development with the Spring Framework[M]. Wrox Press, 2020.',
    '[8] 张宇. 基于JWT的RESTful API认证机制研究[J]. 计算机技术与发展, 2020, 30(3): 120-125.',
    '[9] 赵雷. Redis在Web应用中的缓存策略研究[J]. 计算机工程与设计, 2021, 42(8): 2345-2350.',
    '[10] 周志明. 深入理解Java虚拟机[M]. 北京: 机械工业出版社, 2019.',
]

for ref in refs:
    p = doc.add_paragraph()
    run = p.add_run(ref)
    run.font.name = '宋体'
    run.font.size = Pt(11)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Pt(24)

# ============ 致谢 ============
doc.add_page_break()
add_heading_custom('致  谢', level=1)

add_para('时光荏苒，大学生活即将画上句号。在本毕业设计完成之际，我要向所有给予我帮助和支持的人表示衷心的感谢。')

add_para('首先，我要感谢我的指导老师。在毕业设计的整个过程中，老师从选题、方案设计、技术选型到论文撰写，都给予了耐心细致的指导。老师渊博的专业知识、严谨的治学态度和平易近人的指导方式，让我受益良多。')

add_para('其次，我要感谢我的同学和朋友们。在项目开发过程中，大家相互交流、共同探讨，解决了许多技术难题。这段共同奋斗的经历将成为我宝贵的人生财富。')

add_para('最后，我要感谢我的家人。感谢他们多年来对我学业的支持和鼓励，正是他们的无私付出，让我能够专心完成学业。')

add_para('在此，向所有关心和帮助过我的人致以最诚挚的谢意！')

# ============ 保存文档 ============
output_path = r'D:\develop\毕业设计材料\项目\EduFix_毕业设计说明书.docx'
doc.save(output_path)
print(f'毕业设计文档已生成: {output_path}')
